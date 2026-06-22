import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def add_page_number(run):
    """Insert page number field dynamically into a run."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_callout_box(doc, text, title="Note"):
    """Create a beautiful side-accented note box using an OpenXML-compliant 1x1 table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(5.8)
    
    # Border, shading, and margins in strict schema sequence:
    # 1. tcBorders
    tcBorders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="nil"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="1B365D"/>'
        f'  <w:bottom w:val="nil"/>'
        f'  <w:right w:val="nil"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(parse_xml(tcBorders_xml))
    
    # 2. shd
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))
    
    # 3. tcMar
    tcMar_xml = (
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="120" w:type="dxa"/>'
        f'  <w:bottom w:w="120" w:type="dxa"/>'
        f'  <w:left w:w="180" w:type="dxa"/>'
        f'  <w:right w:w="180" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    cell._tc.get_or_add_tcPr().append(parse_xml(tcMar_xml))
    
    # Add paragraph in cell
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    run_title = p.add_run(f"★ {title}: ")
    run_title.font.name = 'Calibri'
    run_title.font.bold = True
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.italic = True
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    
    # Add spacing after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(2)
    p_after.paragraph_format.space_after = Pt(8)

def add_heading(doc, text, level, space_before=16, space_after=6):
    """Add a styled heading with custom spacing, font, and navy color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri Light'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Primary Dark Navy
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x3B, 0x59, 0x98) # Secondary Steel Blue
    else:
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A) # Dark Slate Gray
    return p

def add_body_paragraph(doc, text, bold_prefix=None, space_after=6, line_spacing=1.15):
    """Add a body paragraph with consistent margins and styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    if bold_prefix:
        brun = p.add_run(bold_prefix)
        brun.font.name = 'Calibri'
        brun.font.size = Pt(11)
        brun.font.bold = True
        brun.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal body
    return p

def add_bullet_item(doc, text, bold_prefix=None):
    """Add a list item with a circular bullet and custom paragraph formatting."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        brun = p.add_run(bold_prefix)
        brun.font.name = 'Calibri'
        brun.font.size = Pt(11)
        brun.font.bold = True
        brun.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_screenshot(doc, image_path, caption):
    """Add an image with a border and caption, centered."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot missing: {image_path}]")
        run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(5.8))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.keep_with_next = True
    
    run_cap = p_cap.add_run(f"Figure: {caption}")
    run_cap.font.name = 'Calibri'
    run_cap.font.size = Pt(9.5)
    run_cap.font.italic = True
    run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_styled_table(doc, headers, rows, widths=None):
    """Create a beautiful table using Word's native 'Medium Shading 1 Accent 1' style, with set column widths."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Apply column widths
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = width
            
    # Format Headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        if len(p.runs) > 0:
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.size = Pt(10)
            
    # Format Body Rows
    for row_idx, data in enumerate(rows):
        row = table.add_row()
        cells = row.cells
        
        # Apply column widths to new row cells
        if widths:
            for i, width in enumerate(widths):
                cells[i].width = width
                
        for i, val in enumerate(data):
            cells[i].text = str(val)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                
    # Add space after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

def generate_report():
    doc = Document()
    
    # ------------------ PAGE SETTINGS ------------------
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Enable different header/footer for the cover page
        section.different_first_page_header_footer = True
        
        # Configure standard page header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("AI Email Agent — Internship Project Report")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        
        # Configure standard page footer (dynamic page numbering)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun = fp.add_run("Page ")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        add_page_number(fp.add_run())

    # ------------------ COVER PAGE ------------------
    # Spacing from top
    for _ in range(2):
        doc.add_paragraph()
        
    # Project Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title_run = p_title.add_run("AI-POWERED AUTOMATED EMAIL AGENT")
    p_title_run.font.name = 'Calibri'
    p_title_run.font.bold = True
    p_title_run.font.size = Pt(28)
    p_title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
    p_title.paragraph_format.space_after = Pt(2)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub_run = p_sub.add_run("Automated Response System Using Anthropic Claude AI, FastAPI, and FastMCP")
    p_sub_run.font.name = 'Calibri Light'
    p_sub_run.font.size = Pt(14)
    p_sub_run.font.color.rgb = RGBColor(0x3B, 0x59, 0x98) # Steel Blue
    p_sub.paragraph_format.space_after = Pt(18)
    
    # Thin blue line under title
    p_div = doc.add_paragraph()
    p_div_run = p_div.add_run("—" * 65)
    p_div_run.font.bold = True
    p_div_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    p_div.paragraph_format.space_after = Pt(40)
    
    # Metadata Title
    p_meta_hdr = doc.add_paragraph()
    p_meta_hdr_run = p_meta_hdr.add_run("INTERNSHIP PROJECT REPORT")
    p_meta_hdr_run.font.name = 'Calibri'
    p_meta_hdr_run.font.bold = True
    p_meta_hdr_run.font.size = Pt(12)
    p_meta_hdr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p_meta_hdr.paragraph_format.space_after = Pt(12)
    
    # Metadata Table (Borderless alignment)
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta_table.autofit = False
    
    # Set columns width
    meta_table.columns[0].width = Inches(2.2)
    meta_table.columns[1].width = Inches(4.0)
    
    metadata_rows = [
        ("Candidate Name:", "Kishan Vadsola"),
        ("Email Address:", "vadsolakishan1310@gmail.com"),
        ("Project Role:", "Software Engineering Intern"),
        ("Host Organization:", "[Internship Company Name]"),
        ("Project Supervisor:", "[Company Supervisor's Name / 'Sir']"),
        ("Submission Date:", "June 22, 2026")
    ]
    
    for idx, (label, val) in enumerate(metadata_rows):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.0)
        
        # Label formatting
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(4)
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Calibri'
        run_lbl.font.bold = True
        run_lbl.font.size = Pt(10.5)
        run_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # Value formatting
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        run_val = p_val.add_run(val)
        run_val.font.name = 'Calibri'
        run_val.font.bold = True
        run_val.font.size = Pt(11)
        run_val.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
    doc.add_page_break()
    
    # ------------------ 1. EXECUTIVE SUMMARY ------------------
    add_heading(doc, "1. Executive Summary", level=1)
    
    add_body_paragraph(doc, 
        "In the modern corporate workspace, email is the primary engine of communication. "
        "However, sorting through high volumes of customer, internal, and promotional messages daily, "
        "drafting contextually accurate replies, and maintaining synchronization leads to severe operational inefficiencies. "
        "The objective of this internship project is to build a secure, intelligent, and autonomous workspace ecosystem. "
        "The resulting application, 'AI Email Agent', combines Large Language Model intelligence with protocol-driven tool architectures "
        "to automatically sync, scan, filter, and draft emails."
    )
    add_body_paragraph(doc,
        "The application consists of three integrated subsystems:\n"
        "1. A robust FastAPI server running concurrent background worker loops that manage Gmail API checks, sync inbox metadata, and cache records in a local database.\n"
        "2. An Anthropic Claude-powered Auto-Reply Engine that uses contextual prompts to reply dynamically according to chosen tones, or safely skip system notifications.\n"
        "3. An HTML5 dashboard frontend incorporating live connection states, telemetry charts, background thread diagnostics, and logs.\n"
        "Additionally, an integration layer using the Model Context Protocol (MCP) exports email operations as tools, making it compatible with agent hosts like Claude Desktop. "
        "The project successfully automates repetitive inbox tasks while preventing reply loop issues."
    )
    
    add_callout_box(doc, 
        "The AI Email Agent is fully functional, utilizing a secure local SQLite cache, Google OAuth2 credentials, and Claude-3.5-Sonnet to provide safe email replies.",
        title="Project Scope"
    )
    
    # ------------------ 2. SYSTEM ARCHITECTURE ------------------
    add_heading(doc, "2. System Architecture & Components", level=1)
    
    add_body_paragraph(doc,
        "The application is structured into a modular system to ensure isolation of concerns, reliable state updates, "
        "and data integrity. The backend handles synchronization protocols and database caches, while the frontend dashboard serves "
        "up-to-date visualization panels."
    )
    
    add_heading(doc, "2.1 Technical Stack", level=2)
    add_bullet_item(doc, "FastAPI (Python): High-performance, async-capable framework used to serve REST APIs and manage application lifespans.", "Backend Server: ")
    add_bullet_item(doc, "Uvicorn: ASGI web server for running the FastAPI application.", "Server runner: ")
    add_bullet_item(doc, "SQLite: Local SQL database used to log sync intervals, configuration settings, inbox records, and auto-reply history.", "Data Store: ")
    add_bullet_item(doc, "Google Gmail API (OAuth 2.0): Authenticates using secure consent, downloads inbox messages, and writes replies.", "Integrations: ")
    add_bullet_item(doc, "Anthropic Claude API: Powers the AI engine that drafts personalized, tone-based email replies.", "AI Brain: ")
    add_bullet_item(doc, "FastMCP: Protocol framework used to create standard MCP endpoints and register tools for Claude Desktop Client.", "MCP Protocol: ")
    add_bullet_item(doc, "Vanilla CSS & JavaScript: Styled dashboard featuring glassmorphism elements, stats telemetry, and detailed previews.", "Frontend: ")
    
    add_heading(doc, "2.2 Module Architecture Diagram", level=2)
    
    add_body_paragraph(doc,
        "The interaction workflow between the client, local server, databases, Google services, and Anthropic APIs is illustrated below:"
    )
    
    # Text-based ascii architecture
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.left_indent = Inches(0.4)
    p_arch.paragraph_format.space_after = Pt(12)
    run_arch = p_arch.add_run(
        "┌────────────────────────────────────────────────────────┐\n"
        "│                    Frontend Dashboard                  │\n"
        "│             (HTML5 / JavaScript / Vanilla CSS)         │\n"
        "└───────────┬─────────────────────────────▲──────────────┘\n"
        "            │ REST Requests               │ SSE / Long-poll updates\n"
        "            ▼                             │\n"
        "┌─────────────────────────────────────────┴──────────────┐\n"
        "│                  FastAPI Backend Server                │\n"
        "│     - Endpoints: /api/emails, /api/settings, etc.      │\n"
        "│     - Background Async Task (Auto-Reply Engine Loop)   │\n"
        "└─────┬──────────────┬─────────────┬─────────────┬───────┘\n"
        "      │ Read/Write   │ OAuth       │ API Call    │ Tool Calls\n"
        "      ▼              ▼             ▼             ▼\n"
        "┌──────────┐   ┌───────────┐ ┌───────────┐ ┌───────────┐\n"
        "│  SQLite  │   │ Gmail API │ │ Claude AI │ │    MCP    │\n"
        "│ Database │   │ (Google)  │ │ (Anthropic│ │  Server   │\n"
        "└──────────┘   └───────────┘ └───────────┘ └───────────┘"
    )
    run_arch.font.name = 'Courier New'
    run_arch.font.size = Pt(9.5)
    
    doc.add_page_break()
    
    # ------------------ 3. KEY FEATURES ------------------
    add_heading(doc, "3. System Features & Implementation Details", level=1)
    
    add_body_paragraph(doc,
        "The project was designed and engineered to include multiple advanced features that ensure automated email execution "
        "is both user-friendly and highly robust. Below are details of the main features implemented:"
    )
    
    add_heading(doc, "3.1 Intelligent Auto-Reply Loop with Label Prevention", level=2)
    add_body_paragraph(doc,
        "Autonomous reply engines are vulnerable to infinite-loop scenarios (where two auto-reply bots continuously reply "
        "to one another). To counter this, the AI Email Agent integrates loop prevention safeguards. When the background "
        "worker thread processes an unread email, it downloads the details, sends them to Claude, drafts the reply, and sends "
        "the response. Crucially, the engine then applies a custom Gmail label ('auto-replied') to the message thread. "
        "During subsequent scan checks, the system queries Gmail for messages that match the criteria 'is:unread' but "
        "specifically excludes threads tagged with the 'auto-replied' label. This blocks duplicate response executions."
    )
    
    add_heading(doc, "3.2 Smart Email Filtering & Filtering Rules", level=2)
    add_body_paragraph(doc,
        "An AI reply should only be generated for active, personal, or corporate inquiries. Automatically responding to "
        "newsletters, commercial flyers, or machine notifications results in spam. The AI Email Agent implements strict "
        "skipping guidelines:"
    )
    add_bullet_item(doc, "Skipping system addresses like no-reply@, noreply@, support@, alert@, news@, notifications@.", "Header Filtering: ")
    add_bullet_item(doc, "Claude checks the mail subject and body text to identify marketing newsletters, sign-up updates, or receipts.", "AI Content Check: ")
    add_bullet_item(doc, "Replies are disabled if the sender is empty, matches automated system templates, or has a mailing list unsubscribe header.", "Safety Defaults: ")
    
    add_heading(doc, "3.3 Multi-Tone Writing Configurations", level=2)
    add_body_paragraph(doc,
        "Depending on the context of the email address, different tones are required. The auto-reply system supports tone "
        "configurations: (1) Professional tone for commercial clients, (2) Friendly tone for colleagues and personal contacts, "
        "and (3) Casual tone for everyday operations. Tone settings can be parsed via the environment variables or toggled "
        "in database preferences."
    )
    
    add_heading(doc, "3.4 HTML5 Live Dashboard UI", level=2)
    add_body_paragraph(doc,
        "A gorgeous glassmorphism dashboard allows the user to supervise the background tasks. It includes: visual cards for "
        "unread count, captured count, and replies sent; a live toggle to turn the background runner ON/OFF; active queue monitoring "
        "showing what the worker is doing; two distinct tab views for mail records and generated AI drafts; search bars to filter "
        "logs; and detailed modal popups that render draft previews."
    )
    
    add_heading(doc, "3.5 Model Context Protocol (MCP) Server Integration", level=2)
    add_body_paragraph(doc,
        "The project includes an MCP server implementation using FastMCP. This wraps the underlying Gmail services "
        "into unified tools: list_unread_emails, get_email_details, reply_to_email, search_emails, and auto_reply_all. "
        "When hooked into an MCP host (such as the Claude Desktop app), it allows desktop LLM interfaces to use "
        "these tools dynamically, creating a decentralized personal agent ecosystem."
    )
    
    doc.add_page_break()
    
    # ------------------ 4. DATABASE DESIGN ------------------
    add_heading(doc, "4. Database Schema & Storage Design", level=1)
    
    add_body_paragraph(doc,
        "To minimize memory footprint while maintaining persistent data logs across server reboots, a local SQLite "
        "database structure was designed. It contains three main tables: settings, email_logs, and inbox_emails."
    )
    
    add_heading(doc, "4.1 SQLite Database Schema Details", level=2)
    
    # Column widths for Schema table
    schema_widths = [Inches(0.9), Inches(1.1), Inches(0.8), Inches(1.2), Inches(2.3)]
    
    # Table definitions
    add_styled_table(doc,
        headers=["Table Name", "Column Name", "Data Type", "Constraints", "Description"],
        rows=[
            ["settings", "key", "TEXT", "PRIMARY KEY", "Unique configuration keyword (e.g. is_auto_reply_on)"],
            ["settings", "value", "TEXT", "", "Assigned configuration setting value"],
            
            ["email_logs", "id", "INTEGER", "PRIMARY KEY AUTOINCREMENT", "Unique internal transaction ID"],
            ["email_logs", "message_id", "TEXT", "UNIQUE", "Gmail API unique string identifier"],
            ["email_logs", "sender", "TEXT", "", "Email address of the sender replied to"],
            ["email_logs", "subject", "TEXT", "", "Subject line of the email thread"],
            ["email_logs", "reply_body", "TEXT", "", "The AI draft body sent as response"],
            ["email_logs", "timestamp", "TEXT", "", "ISO 8601 string representing transaction time"],
            
            ["inbox_emails", "id", "INTEGER", "PRIMARY KEY AUTOINCREMENT", "Unique sync counter ID"],
            ["inbox_emails", "message_id", "TEXT", "UNIQUE", "Gmail API message ID used to prevent duplicate entry"],
            ["inbox_emails", "sender", "TEXT", "", "Sender email address"],
            ["inbox_emails", "sender_name", "TEXT", "", "Human readable name of the sender"],
            ["inbox_emails", "subject", "TEXT", "", "Subject line of the inbox message"],
            ["inbox_emails", "snippet", "TEXT", "", "Truncated snippet text preview"],
            ["inbox_emails", "date", "TEXT", "", "Origin date header of the email"],
            ["inbox_emails", "captured_at", "TEXT", "", "Local database capture timestamp"],
            ["inbox_emails", "is_replied", "INTEGER", "DEFAULT 0", "Flag indicator (1 = Auto-replied, 0 = Outstanding)"]
        ],
        widths=schema_widths
    )
    
    add_heading(doc, "4.2 Sync & Log Workflow", level=2)
    add_body_paragraph(doc,
        "Every 30 seconds, the background loop executes a synchronization block. It scrapes the user's unread folders, "
        "and commits each new message entry to the `inbox_emails` table. If the auto-reply engine toggle is active, it runs "
        "the response engine. Upon a successful response API call to Gmail, the agent writes a log to the `email_logs` "
        "table and updates the status flag `is_replied = 1` in the `inbox_emails` table. The frontend polls these SQLite tables "
        "to serve up-to-date metrics."
    )
    
    doc.add_page_break()
    
    # ------------------ 5. API DOCUMENTATION ------------------
    add_heading(doc, "5. REST API Documentation", level=1)
    
    add_body_paragraph(doc,
        "The backend exposes REST API endpoints, allowing the dashboard frontend or outer integrations to manage "
        "the email engine. All routes are prefixed with '/api'."
    )
    
    # Column widths for API table
    api_widths = [Inches(0.7), Inches(1.5), Inches(1.5), Inches(1.0), Inches(1.6)]
    
    add_styled_table(doc,
        headers=["HTTP Method", "Endpoint Route", "Description", "Parameters", "Response Schema"],
        rows=[
            ["GET", "/api/health", "Service status check", "None", '{"status": "ok", "service": "on/off"}'],
            ["GET", "/api/emails/unread", "Retrieve unread emails", "None", "List of pending unread emails"],
            ["GET", "/api/emails/{id}", "Fetch email details", "id (path parameter)", "Email subject, body, snippet"],
            ["POST", "/api/emails/{id}/reply", "Trigger manual AI reply", "id (path), body (json optional)", '{"status": "replied", "body": "..."}'],
            ["POST", "/api/emails/auto-reply", "Trigger full auto-reply scan", "None", '{"status": "completed", "processed": N}'],
            ["GET", "/api/emails/search", "Search email logs", "q (query string search)", "List of matching database rows"],
            ["POST", "/api/settings/toggle", "Toggle auto-reply engine", "is_on (boolean body)", '{"status": "success", "is_on": boolean}']
        ],
        widths=api_widths
    )
    
    # ------------------ 6. SYSTEM SCREENSHOTS ------------------
    doc.add_page_break()
    add_heading(doc, "6. Working Implementation Screenshots", level=1)
    
    add_body_paragraph(doc,
        "This section features screenshots captured from the live application showing the frontend web control dashboard, "
        "the AI response history logs, and the email metadata modal popup interface."
    )
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    img_dashboard_inbox = os.path.join(script_dir, "screenshots", "dashboard_inbox.png")
    img_dashboard_replies = os.path.join(script_dir, "screenshots", "dashboard_replies.png")
    img_detail_modal = os.path.join(script_dir, "screenshots", "detail_modal.png")
    
    add_heading(doc, "6.1 Main Dashboard & Inbox View", level=2)
    add_body_paragraph(doc,
        "The primary view displays the system status, email counters, and the live control toggle. Under the 'Mail Inbox' "
        "tab, users can inspect all captured inbox emails sorted in reverse chronological order."
    )
    add_screenshot(doc, img_dashboard_inbox, "AI Email Agent Dashboard — Mail Inbox view showing live indicators, statistics row, and database sync.")
    
    add_heading(doc, "6.2 AI Replies Log View", level=2)
    add_body_paragraph(doc,
        "Selecting the 'AI Replies' tab switches the bottom workspace list to the list of successfully sent auto-replies, "
        "complete with details about the receiver, subject lines, and the exact timestamps of execution."
    )
    add_screenshot(doc, img_dashboard_replies, "AI Email Agent Dashboard — AI Replies Log history view tracking sent email logs.")
    
    add_heading(doc, "6.3 AI Draft Preview Modal Interface", level=2)
    add_body_paragraph(doc,
        "Clicking on any log entry displays a modal window overlay. This displays the full generated AI response body, "
        "allowing the operator to read what the agent drafted for that email."
    )
    add_screenshot(doc, img_detail_modal, "AI Email Agent Dashboard — Detail Modal rendering full AI response draft and email headers.")
    
    doc.add_page_break()
    
    # ------------------ 7. CHALLENGES & SOLUTIONS ------------------
    add_heading(doc, "7. Technical Challenges & Solutions", level=1)
    
    add_body_paragraph(doc, "During the implementation, several engineering challenges were encountered and solved:")
    
    add_body_paragraph(doc,
        "Connecting to Gmail API from local development servers requires Google OAuth consent. Since the refresh tokens "
        "expire if left inactive or if permissions change, we designed a file-cached token flow (`token.json`). "
        "On first launch, if no token is found, a browser consent screen opens. The user approves, and the server "
        "saves a serialized token. Subsequent calls use this token, refreshing it automatically if expired.",
        bold_prefix="Challenge 1: Google OAuth Caching & Re-authorization. "
    )
    
    add_body_paragraph(doc,
        "Integrating a dashboard with dynamic content requires background updates. If frontend code polls the backend "
        "continuously, it creates unnecessary network traffic. We optimized this by incorporating an SQLite-backed status cache. "
        "The background worker thread updates database entries, and the frontend queries a single light status endpoint "
        "to decide if it should re-fetch mail items, minimizing resource utilization.",
        bold_prefix="Challenge 2: State Sync & Performance Optimization. "
    )
    
    add_body_paragraph(doc,
        "If a thread is locked due to an API timeout, it can crash the web dashboard. We solved this by using "
        "asyncio tasks with cooperative multi-tasking. The background worker is mounted as an independent async lifespan "
        "task that runs alongside FastAPI, allowing non-blocking concurrent request routing.",
        bold_prefix="Challenge 3: Background Worker Reliability & Thread Concurrency. "
    )
    
    # ------------------ 8. FUTURE ENHANCEMENTS ------------------
    add_heading(doc, "8. Future Enhancements & Extensions", level=1)
    
    add_bullet_item(doc, "Add a draft review step where replies are held in a pending queue until approved by a human agent.", "Human-in-the-Loop Review: ")
    add_bullet_item(doc, "Provide options to customize custom rules, skip words, and specific responder tones per email folder/label.", "Granular Settings: ")
    add_bullet_item(doc, "Implement support for IMAP/SMTP and Outlook Graph API, opening the platform to users outside the Google ecosystem.", "Multi-Provider Mail Support: ")
    add_bullet_item(doc, "Utilize local LLMs (via Ollama or Llama.cpp) to run email processing fully offline, lowering API usage expenses and protecting private data.", "Local LLM Integration: ")
    
    # ------------------ 9. CONCLUSION ------------------
    add_heading(doc, "9. Conclusion & Experience Gained", level=1)
    add_body_paragraph(doc,
        "This internship project has successfully demonstrated the viability of using LLMs and modern protocol architectures "
        "to automate office productivity tasks. The AI Email Agent is fully functional, secure, and ready for deployment. "
        "Through this engineering process, I gained deep hands-on experience in API integrations, OAuth authentication, SQLite "
        "database schemas, background thread workers in python web servers, and building elegant modern CSS/JS interfaces."
    )
    
    # Save the document
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "AI_Email_Agent_Internship_Report_Final.docx")
    doc.save(out_path)
    print(f"Report successfully generated at: {out_path}")

if __name__ == "__main__":
    generate_report()
